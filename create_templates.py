"""
Creates the two invoice Word templates required by invoice_service.py.
Run once: python create_templates.py
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _set_cell_bg(cell, hex_color: str):
    """Set cell background colour via raw XML (e.g. 'BDD7EE')."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_col_width(table, col_index: int, width_cm: float):
    """Set preferred width of every cell in a column."""
    from docx.shared import Cm as _Cm
    for row in table.rows:
        cell = row.cells[col_index]
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = OxmlElement("w:tcW")
        tcW.set(qn("w:w"), str(int(_Cm(width_cm).emu / 914.4 * 20)))  # twips
        tcW.set(qn("w:type"), "dxa")
        tcPr.append(tcW)


def _header_block(doc: Document):
    """Add the shared company-placeholder header to both templates."""
    p = doc.add_paragraph()
    run = p.add_run("[Firmenname] · [Straße Nr.] · [PLZ Ort] · [Tel/E-Mail]")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    p.paragraph_format.space_after = Pt(2)

    # thin separator line
    p2 = doc.add_paragraph()
    p2.add_run("─" * 80).font.size = Pt(8)
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(6)


def _invoice_meta(doc: Document):
    """Rechnungsnummer + Datum block."""
    p = doc.add_paragraph()
    p.add_run("Rechnungsnummer:  ").bold = True
    p.add_run("{{invoice_number}}")

    p = doc.add_paragraph()
    p.add_run("Datum:  ").bold = True
    p.add_run("{{invoice_date}}")

    doc.add_paragraph()


def _recipient_block(doc: Document):
    """An / supplier block."""
    p = doc.add_paragraph()
    p.add_run("An:").bold = True

    doc.add_paragraph("{{supplier_name}}")
    doc.add_paragraph("{{supplier_contact}}")
    doc.add_paragraph()


def _subject_block(doc: Document):
    """Betreff + Zeitraum."""
    p = doc.add_paragraph()
    p.add_run("Betreff:  ").bold = True
    p.add_run("{{entry_type}} \u2013 {{description}}")

    p = doc.add_paragraph()
    p.add_run("Zeitraum:  ").bold = True
    p.add_run("{{period_start}} bis {{period_end}}")

    doc.add_paragraph()


def _totals_block(doc: Document):
    """Nettobetrag / MwSt / Gesamt."""
    p = doc.add_paragraph()
    p.add_run("Nettobetrag:  ").bold = True
    p.add_run("{{net_amount}} \u20ac")

    p = doc.add_paragraph()
    p.add_run("MwSt. {{tax_rate}} %:  ").bold = True
    p.add_run("{{tax_amount}} \u20ac")

    p = doc.add_paragraph()
    run = p.add_run("Gesamtbetrag:  ")
    run.bold = True
    run.font.size = Pt(13)
    run2 = p.add_run("{{total_amount}} \u20ac")
    run2.bold = True
    run2.font.size = Pt(13)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("Notizen:  ").bold = True
    p.add_run("{{notes}}")


# ---------------------------------------------------------------------------
# Standard template  (WKZ fixed, WKZ %, Umsatzbonus)
# ---------------------------------------------------------------------------

def create_standard_template(path: str):
    doc = Document()

    # Margins
    sec = doc.sections[0]
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)

    _header_block(doc)
    _invoice_meta(doc)
    _recipient_block(doc)
    _subject_block(doc)

    # Description line
    p = doc.add_paragraph()
    p.add_run("Beschreibung:  ").bold = True
    p.add_run("{{line_description}}")
    doc.add_paragraph()

    # --- Conditional block: WKZ percentage ---
    # {%p if base_amount %} / content / {%p endif %}  — docxtpl removes the control paragraphs
    doc.add_paragraph("{%p if base_amount %}")
    p = doc.add_paragraph()
    p.add_run("Basis (Einkaufsumsatz):  ").bold = True
    p.add_run("{{base_amount}} \u20ac")
    doc.add_paragraph("{%p endif %}")

    doc.add_paragraph("{%p if percentage %}")
    p = doc.add_paragraph()
    p.add_run("Prozentsatz:  ").bold = True
    p.add_run("{{percentage}} %")
    doc.add_paragraph("{%p endif %}")

    # --- Conditional block: Umsatzbonus ---
    doc.add_paragraph("{%p if achieved_revenue %}")
    p = doc.add_paragraph()
    p.add_run("Erzielter Umsatz:  ").bold = True
    p.add_run("{{achieved_revenue}} \u20ac")
    doc.add_paragraph("{%p endif %}")

    doc.add_paragraph("{%p if bonus_tier %}")
    p = doc.add_paragraph()
    p.add_run("Staffel ab:  ").bold = True
    p.add_run("{{bonus_tier}} \u20ac")
    doc.add_paragraph("{%p endif %}")

    doc.add_paragraph("{%p if bonus_percentage %}")
    p = doc.add_paragraph()
    p.add_run("Bonussatz:  ").bold = True
    p.add_run("{{bonus_percentage}} %")
    doc.add_paragraph("{%p endif %}")

    doc.add_paragraph()
    _totals_block(doc)

    os.makedirs(os.path.dirname(path), exist_ok=True)
    doc.save(path)
    print(f"  Created: {path}")


# ---------------------------------------------------------------------------
# Kickback template  (article table with row-repeat)
# ---------------------------------------------------------------------------

def create_kickback_template(path: str):
    doc = Document()

    sec = doc.sections[0]
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)

    _header_block(doc)
    _invoice_meta(doc)
    _recipient_block(doc)
    _subject_block(doc)

    # Table: header | loop-start | data | loop-end  (4 rows × 5 cols)
    table = doc.add_table(rows=4, cols=5)
    table.style = "Table Grid"

    # --- Row 0: Header ---
    hdr = table.rows[0]
    for i, text in enumerate(["Art.-Nr.", "Beschreibung", "Menge", "Satz (\u20ac)", "Betrag (\u20ac)"]):
        cell = hdr.cells[i]
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        _set_cell_bg(cell, "BDD7EE")

    # --- Row 1: Loop start marker ---
    loop_start = table.rows[1]
    loop_start.cells[0].paragraphs[0].clear()
    loop_start.cells[0].paragraphs[0].add_run("{%tr for item in items %}")
    for i in range(1, 5):
        loop_start.cells[i].paragraphs[0].clear()

    # --- Row 2: Data row (gets repeated per item) ---
    data_row = table.rows[2]
    placeholders = [
        "{{item.article_number}}",
        "",
        "{{item.qty}}",
        "{{item.rate}}",
        "{{item.line_amount}}",
    ]
    for i, text in enumerate(placeholders):
        data_row.cells[i].paragraphs[0].clear()
        if text:
            data_row.cells[i].paragraphs[0].add_run(text)

    # --- Row 3: Loop end marker ---
    loop_end = table.rows[3]
    loop_end.cells[0].paragraphs[0].clear()
    loop_end.cells[0].paragraphs[0].add_run("{%tr endfor %}")
    for i in range(1, 5):
        loop_end.cells[i].paragraphs[0].clear()

    # Column widths (twips: 1 cm ≈ 567 twips)
    col_widths_cm = [2.8, 5.0, 2.0, 2.5, 2.5]
    for col_i, w in enumerate(col_widths_cm):
        _set_col_width(table, col_i, w)

    doc.add_paragraph()
    _totals_block(doc)

    os.makedirs(os.path.dirname(path), exist_ok=True)
    doc.save(path)
    print(f"  Created: {path}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    base = os.path.join(os.path.dirname(__file__), "data", "templates")
    print("Generating invoice templates...")
    create_standard_template(os.path.join(base, "invoice_standard.docx"))
    create_kickback_template(os.path.join(base, "invoice_kickback.docx"))
    print("Done.")
